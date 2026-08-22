import io
import re
from typing import Dict, Any

KNOWN_ROLES = [
    'Full Stack Developer', 'Full Stack Engineer',
    'Frontend Developer', 'Frontend Engineer', 'UI Developer',
    'Backend Developer', 'Backend Engineer',
    'Machine Learning Engineer', 'AI Engineer', 'NLP Engineer', 'Deep Learning Engineer',
    'Data Scientist', 'Data Analyst', 'Data Engineer',
    'DevOps Engineer', 'Cloud Engineer', 'Site Reliability Engineer',
    'Software Engineer', 'Software Developer',
    'QA Engineer', 'Mobile Developer', 'Android Developer', 'iOS Developer',
    'HR Specialist', 'HR Manager', 'Talent Acquisition Specialist',
    'Product Manager', 'Project Manager', 'UI/UX Designer'
]

def detect_role_from_text(clean_text: str) -> str:
    """Detect candidate's professional title/domain from resume text."""
    lines = [line.strip() for line in clean_text.split('\n') if line.strip()]
    
    # 1. Check first 6 lines
    for line in lines[:6]:
        for role in KNOWN_ROLES:
            if re.search(r'\b' + re.escape(role) + r'\b', line, re.IGNORECASE):
                return role
                
    # 2. Check summary / objective section
    summary_match = re.search(r'(?:summary|profile|about me|objective)[:\s\n]+([^\n\.]+)', clean_text, re.IGNORECASE)
    if summary_match:
        summary_line = summary_match.group(1)
        for role in KNOWN_ROLES:
            if re.search(r'\b' + re.escape(role) + r'\b', summary_line, re.IGNORECASE):
                return role

    # 3. Keyword-based heuristic fallback
    text_lower = clean_text.lower()
    if 'full stack' in text_lower or ('react' in text_lower and ('node' in text_lower or 'express' in text_lower or 'django' in text_lower)):
        return 'Full Stack Developer'
    if 'frontend' in text_lower or 'react' in text_lower or 'vue' in text_lower:
        return 'Frontend Developer'
    if 'backend' in text_lower or 'fastapi' in text_lower or 'django' in text_lower or 'spring boot' in text_lower:
        return 'Backend Developer'
    if 'devops' in text_lower or 'kubernetes' in text_lower or 'docker' in text_lower or 'terraform' in text_lower:
        return 'DevOps Engineer'
    if 'machine learning' in text_lower or 'pytorch' in text_lower or 'deep learning' in text_lower:
        return 'Machine Learning Engineer'
    if 'data scientist' in text_lower:
        return 'Data Scientist'
    if 'data analyst' in text_lower or 'tableau' in text_lower or 'power bi' in text_lower:
        return 'Data Analyst'
    if 'hr' in text_lower or 'human resource' in text_lower or 'recruitment' in text_lower:
        return 'HR Specialist'
        
    return 'Software Engineer'

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"  [WARN] PDF extraction error: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        return "\n".join(full_text).strip()
    except Exception as e:
        print(f"  [WARN] DOCX extraction error: {e}")
        return ""

def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode('utf-8').strip()
    except UnicodeDecodeError:
        return file_bytes.decode('latin-1', errors='ignore').strip()

def parse_resume_file(filename: str, file_bytes: bytes) -> Dict[str, Any]:
    ext = filename.lower().split('.')[-1]
    
    if ext == 'pdf':
        text = extract_text_from_pdf(file_bytes)
    elif ext in ['docx', 'doc']:
        text = extract_text_from_docx(file_bytes)
    else:
        text = extract_text_from_txt(file_bytes)

    clean_text = re.sub(r'[ \t]+', ' ', text)
    clean_text = re.sub(r'\n{3,}', '\n\n', clean_text).strip()
    
    words = re.findall(r'\w+', clean_text)
    word_count = len(words)

    candidate_name = "Candidate"
    lines = [line.strip() for line in clean_text.split('\n') if line.strip()]
    for line in lines[:3]:
        if not re.search(r'resume|curriculum|vitae|summary|experience|education|contact|phone|email|profile', line, re.IGNORECASE):
            if len(line.split()) <= 4 and len(line) <= 40:
                candidate_name = line.replace('|', '').strip()
                break

    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', clean_text)
    email = email_match.group(0) if email_match else None
    
    detected_role = detect_role_from_text(clean_text)

    return {
        "status": "success" if word_count > 15 else "warning",
        "filename": filename,
        "candidate_name": candidate_name,
        "detected_role": detected_role,
        "email": email,
        "word_count": word_count,
        "extracted_text": clean_text
    }
